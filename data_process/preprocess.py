import os
from docx import Document

from pandas import DataFrame

# docx这个库只能处理docx后缀文件，不能处理doc后缀，用重命名的方法不行。
def rename_filename_from_doc_to_docx(folder):
    for filename in os.listdir(folder):
        if filename.endswith('.doc'):
            os.rename(os.path.join(folder, filename), os.path.join(folder, filename.split('.')[0] + '.docx'))

def docx_to_onetxt(folder_path, output_txt_path):
    # 创建一个空的列表来存储所有文本内容
    all_text = []

    # 遍历文件夹中的所有文件
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            # 构建文件的完整路径
            file_path = os.path.join(folder_path, filename)
            print(file_path)
            # 打开.docx文件并读取文本内容
            doc = Document(file_path)
            # 一篇doc文档
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            old_content = ''.join(full_text)
            new_content = ''
            for x in old_content:
                if x in ['', ' ', '\n', ' ','\u3000']:
                    continue
                new_content += x
            all_text.append(new_content)
    print(all_text)

    # 将所有文本内容写入新的.txt文件
    with open(output_txt_path, 'w', encoding='utf-8') as txt_file:
        txt_file.writelines(all_text)

def docx_to_txt(input_folder_path, output_folder_path):
    if not os.path.exists(output_folder_path):
        os.makedirs(output_folder_path)
    # 遍历文件夹中的所有文件
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            # 构建文件的完整路径
            docx_path = os.path.join(input_folder_path, filename)
            txt_path = os.path.join(output_folder_path, filename.split('.')[0]+'.txt')
            print(docx_path)
            # 打开.docx文件并读取文本内容
            doc = Document(docx_path)
            # 一篇doc文档
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            old_content = ''.join(full_text)
            new_content = ''
            for x in old_content:
                if x in ['', ' ', '\n', ' ', '\u3000']:
                    continue
                new_content += x
            # 将所有文本内容写入新的.txt文件
            with open(txt_path, 'w', encoding='utf-8') as txt_file:
                txt_file.writelines(new_content)

# 单个txt文本数据预处理
def solve_data(input_path,output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        with open(output_path, 'w', encoding='utf-8') as out:
            ss = ''
            for line in f.readlines():
                for x in line:
                    if x in ['',' ','\n',' ','\u3000','\u3000']:
                        continue
                    ss += x
            out.write(ss)
            out.flush()


if __name__ == '__main__':
    # input_file_path = "D:\\APersonal\\Agraduate\\AinGroup\\AProjects\\NER\\data_process\\onedata.txt"
    # output_file_path = "D:\\APersonal\\Agraduate\\AinGroup\\AProjects\\NER\\data_process\\onedata_output.txt"
    # solve_data(input_file_path,output_file_path)

    # rename_filename_from_doc_to_docx('D:\\APersonal\\Agraduate\\AinGroup\\AProjects\\NER\\data_process\\爆炸')

    folder_path = 'D:\\APersonal\\Agraduate\\AinGroup\\AProjects\\NER\\data_process\\爆炸_docx'  # 替换为你的文件夹路径
    output_txt_path = 'D:\\APersonal\\Agraduate\\AinGroup\\AProjects\\NER\\data_process\\爆炸2.txt'  # 输出的txt文件名
    output_floder_path = 'D:\\APersonal\\Agraduate\\AinGroup\\AProjects\\NER\\data_process\\爆炸_txt'
    # docx_to_onetxt(folder_path, output_txt_path)
    docx_to_txt(folder_path, output_floder_path)

    # # 中文全角空格
    # s = '\u3000'
    # print(len(s))

